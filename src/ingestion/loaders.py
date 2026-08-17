"""
Multi-format Document Loaders for Production RAG Engine.
Supports PDF (.pdf), Microsoft Word (.docx), Plain Text (.txt), and Markdown (.md).
"""
import os
from pathlib import Path
from typing import List, Union
from langchain_core.documents import Document
from src.core.logger import get_logger

logger = get_logger("document_loader")


class DocumentLoader:
    """
    Robust multi-format document loader with rich metadata extraction.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

    @classmethod
    def load_file(cls, file_path: Union[str, Path]) -> List[Document]:
        """
        Load a single document from file path and return standard LangChain Document objects.
        """
        path = Path(file_path).resolve()

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        ext = path.suffix.lower()
        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file format '{ext}'. Supported formats: {', '.join(cls.SUPPORTED_EXTENSIONS)}"
            )

        logger.info(f"Loading document: [bold green]{path.name}[/bold green] ({ext})")

        try:
            if ext == ".pdf":
                return cls._load_pdf(path)
            elif ext == ".docx":
                return cls._load_docx(path)
            elif ext in {".txt", ".md"}:
                return cls._load_text(path)
            else:
                return []
        except Exception as e:
            logger.error(f"Failed to load document {path.name}: {str(e)}")
            raise e

    @classmethod
    def load_directory(cls, directory_path: Union[str, Path], recursive: bool = True) -> List[Document]:
        """
        Load all supported documents from a directory.
        """
        dir_path = Path(directory_path).resolve()
        if not dir_path.exists() or not dir_path.is_dir():
            raise NotADirectoryError(f"Directory not found: {dir_path}")

        pattern = "**/*" if recursive else "*"
        all_docs: List[Document] = []
        files = [p for p in dir_path.glob(pattern) if p.is_file() and p.suffix.lower() in cls.SUPPORTED_EXTENSIONS]

        logger.info(f"Found {len(files)} supported document(s) in {dir_path}")

        for file_path in files:
            try:
                docs = cls.load_file(file_path)
                all_docs.extend(docs)
            except Exception as err:
                logger.warning(f"Skipping corrupt or unreadable file: {file_path.name} ({err})")

        logger.info(f"Successfully loaded {len(all_docs)} total document page(s)/section(s)")
        return all_docs

    @staticmethod
    def _load_pdf(path: Path) -> List[Document]:
        """Extract text and page-level metadata from PDF using pypdf."""
        import pypdf

        documents: List[Document] = []
        file_size = path.stat().st_size

        with open(path, "rb") as f:
            reader = pypdf.PdfReader(f)
            total_pages = len(reader.pages)

            for page_idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                cleaned_text = text.strip()
                if not cleaned_text:
                    continue

                metadata = {
                    "source": str(path),
                    "file_name": path.name,
                    "file_type": "pdf",
                    "file_size_bytes": file_size,
                    "page_number": page_idx + 1,
                    "total_pages": total_pages,
                    "char_count": len(cleaned_text),
                }
                documents.append(Document(page_content=cleaned_text, metadata=metadata))

        return documents

    @staticmethod
    def _load_docx(path: Path) -> List[Document]:
        """Extract text from Word Document (.docx) using python-docx."""
        import docx

        doc = docx.Document(path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        if not full_text:
            return []

        metadata = {
            "source": str(path),
            "file_name": path.name,
            "file_type": "docx",
            "file_size_bytes": path.stat().st_size,
            "page_number": 1,
            "total_pages": 1,
            "char_count": len(full_text),
        }
        return [Document(page_content=full_text, metadata=metadata)]

    @staticmethod
    def _load_text(path: Path) -> List[Document]:
        """Extract text from .txt and .md files."""
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read().strip()

        if not content:
            return []

        metadata = {
            "source": str(path),
            "file_name": path.name,
            "file_type": path.suffix.lower().lstrip("."),
            "file_size_bytes": path.stat().st_size,
            "page_number": 1,
            "total_pages": 1,
            "char_count": len(content),
        }
        return [Document(page_content=content, metadata=metadata)]
